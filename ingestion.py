"""Loads and indexes documents into ChromaDB."""

from __future__ import annotations

import json
import re
import os
import shutil
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable
from urllib.parse import urljoin, urlparse

import chromadb
import requests
from bs4 import BeautifulSoup
from llama_index.core import Document, StorageContext, VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore

from source_loader import PREPARED_SOURCE_METADATA, clear_prepared_source, prepare_sources
import source_loader


PROJECT_ROOT = Path(__file__).resolve().parent
RAW_DIR = PROJECT_ROOT / "data" / "raw"
CHROMA_DIR = PROJECT_ROOT / "chroma_db"
CHROMA_COLLECTION_NAME = "advanced_rag"
PYTHON_TUTORIAL_URL = "https://docs.python.org/3/tutorial/"
PAGE_LIMIT = 50
CHUNK_SIZE_TOKENS = 512
CHUNK_OVERLAP_TOKENS = 64
EMBED_MODEL = os.getenv("EMBED_MODEL", "BAAI/bge-small-en-v1.5")
CHUNK_ABLATION_SIZES = [256, 512, 1024]
EMBEDDING_COMPARISON_MODELS = [
    "BAAI/bge-small-en-v1.5",
    "BAAI/bge-base-en-v1.5",
    "intfloat/e5-small-v2",
]
SOURCE_EXTENSIONS = {
    ".c",
    ".cfg",
    ".cs",
    ".cpp",
    ".cxx",
    ".docx",
    ".go",
    ".h",
    ".hpp",
    ".ini",
    ".java",
    ".js",
    ".json",
    ".jsx",
    ".kt",
    ".md",
    ".mjs",
    ".pdf",
    ".py",
    ".pyi",
    ".rb",
    ".rs",
    ".rst",
    ".sh",
    ".sql",
    ".toml",
    ".ts",
    ".tsx",
    ".txt",
    ".yaml",
    ".yml",
}
SOURCE_FILENAMES = {"Dockerfile", "Makefile", "Procfile"}
IGNORED_DIR_NAMES = {
    ".git",
    ".hg",
    ".idea",
    ".svn",
    ".venv",
    ".vscode",
    "build",
    "chroma_db",
    "coverage",
    "dist",
    "env",
    "node_modules",
    "target",
    "venv",
    "__pycache__",
}

CURRENT_SOURCE_PATH = PROJECT_ROOT / "data" / "current_source.json"


def _detect_source_type(raw_dir: Path) -> str:
    """Inspect files in raw_dir to detect 'local', 'github', or 'huggingface'."""
    dir_name = raw_dir.name.lower()
    if "huggingface" in dir_name or "hf" in dir_name:
        return "huggingface"
    if "github" in dir_name or "gh" in dir_name:
        return "github"
    return "local"


def _source_slug_from_raw(raw_dir: Path) -> str:
    """Extract the slug from the raw_dir subdirectory name (first-level under data/raw/)."""
    return raw_dir.name


def write_current_source(
    raw_dir: Path,
    file_count: int,
    chunk_count: int,
    source_input: str = "",
    source_type: str | None = None,
    source_slug: str | None = None,
) -> None:
    """Write data/current_source.json with metadata about the indexed source."""
    data = {
        "source_input": source_input,
        "source_type": source_type or _detect_source_type(raw_dir),
        "source_slug": source_slug or _source_slug_from_raw(raw_dir),
        "indexed_at": datetime.now(timezone.utc).isoformat(),
        "file_count": file_count,
        "chunk_count": chunk_count,
    }
    CURRENT_SOURCE_PATH.parent.mkdir(parents=True, exist_ok=True)
    CURRENT_SOURCE_PATH.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def load_current_source() -> dict | None:
    """Load data/current_source.json; return None if missing."""
    if not CURRENT_SOURCE_PATH.exists():
        return None
    return json.loads(CURRENT_SOURCE_PATH.read_text(encoding="utf-8"))


def _safe_filename(url: str) -> str:
    path = urlparse(url).path.strip("/")
    filename = re.sub(r"[^A-Za-z0-9_.-]+", "_", path).strip("_")
    return f"{filename or 'python_tutorial_index'}.txt"


def _extract_pdf_text(path: Path) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "PDF ingestion requires PyPDF2. Install the project requirements before indexing .pdf sources."
        ) from exc

    reader = PdfReader(str(path))
    page_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            page_text.append(text.strip())
    return "\n\n".join(page_text)


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError(
            "DOCX ingestion requires python-docx. Install the project requirements before indexing .docx sources."
        ) from exc

    document = DocxDocument(str(path))
    text_parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append("\t".join(cells))
    return "\n".join(text_parts)


def _safe_read_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="ignore")


def _is_supported_source_file(path: Path) -> bool:
    return path.is_file() and (
        path.suffix.lower() in SOURCE_EXTENSIONS or path.name in SOURCE_FILENAMES
    ) and not any(part in IGNORED_DIR_NAMES for part in path.parts)


def _extract_page_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "nav", "footer"]):
        tag.decompose()

    title = soup.find("h1")
    body = soup.find("div", class_="body") or soup.find("main") or soup.body or soup
    lines = [line.strip() for line in body.get_text("\n").splitlines()]
    text = "\n".join(line for line in lines if line)

    if title and title.get_text(strip=True) not in text[:200]:
        text = f"{title.get_text(strip=True)}\n\n{text}"
    return text


def _discover_tutorial_pages(session: requests.Session, limit: int = PAGE_LIMIT) -> list[str]:
    pending = [PYTHON_TUTORIAL_URL]
    visited: set[str] = set()
    pages: list[str] = []

    while pending and len(pages) < limit:
        url = pending.pop(0)
        if url in visited:
            continue
        visited.add(url)

        response = session.get(url, timeout=20)
        response.raise_for_status()
        pages.append(url)

        soup = BeautifulSoup(response.text, "html.parser")
        for anchor in soup.select("a[href]"):
            next_url = urljoin(url, anchor["href"]).split("#", 1)[0]
            parsed = urlparse(next_url)
            if parsed.netloc != "docs.python.org":
                continue
            if not parsed.path.startswith("/3/tutorial/"):
                continue
            if parsed.path.endswith("/") or parsed.path.endswith(".html"):
                if next_url not in visited and next_url not in pending:
                    pending.append(next_url)

    return pages


def download_python_tutorial_pages(raw_dir: Path = RAW_DIR, limit: int = PAGE_LIMIT) -> list[Path]:
    """Download Python tutorial pages as plain text files."""

    raw_dir.mkdir(parents=True, exist_ok=True)
    session = requests.Session()
    session.headers.update({"User-Agent": "advanced-rag-ingestion/1.0"})

    saved_files: list[Path] = []
    for url in _discover_tutorial_pages(session, limit=limit):
        response = session.get(url, timeout=20)
        response.raise_for_status()
        text = _extract_page_text(response.text)
        target = raw_dir / _safe_filename(url)
        target.write_text(f"Source: {url}\n\n{text}\n", encoding="utf-8")
        saved_files.append(target)

    return saved_files


def _source_files(raw_dir: Path = RAW_DIR) -> list[Path]:
    if not raw_dir.exists():
        return []

    return sorted(path for path in raw_dir.rglob("*") if _is_supported_source_file(path))


def _current_source_raw_dir(raw_dir: Path, current_source: dict | None) -> Path | None:
    if not current_source:
        return None
    source_slug = current_source.get("source_slug")
    if not source_slug or not raw_dir.exists():
        return None
    for child in raw_dir.iterdir():
        if child.is_dir() and child.name.lower() == str(source_slug).lower():
            return child
    return None


def load_or_download_sources(raw_dir: Path = RAW_DIR, limit: int = PAGE_LIMIT) -> list[Path]:
    """Use local text/markdown sources when present; otherwise scrape Python docs."""

    raw_dir.mkdir(parents=True, exist_ok=True)
    existing = _source_files(raw_dir)
    if existing:
        return existing
    if os.getenv("ALLOW_DOCS_DOWNLOAD") != "1":
        raise RuntimeError(
            "No source files found in data/raw. Add repository files with prepare_sources(), "
            "or set ALLOW_DOCS_DOWNLOAD=1 to download Python tutorial docs explicitly."
        )
    return download_python_tutorial_pages(raw_dir=raw_dir, limit=limit)


def _metadata_for_files(
    files: list[Path],
    current_source: dict | None,
    source_files_were_explicit: bool,
    raw_dir: Path,
) -> dict | None:
    for path in files:
        metadata = PREPARED_SOURCE_METADATA.get(path.resolve().as_posix())
        if metadata:
            return metadata
    prepared_source = source_loader.load_prepared_source(source_loader.prepared_source_path_for_raw_dir(raw_dir))
    if prepared_source:
        prepared_root = raw_dir / str(prepared_source.get("source_slug", ""))
        resolved_root = prepared_root.resolve(strict=False)
        if any(_path_is_relative_to(path.resolve(strict=False), resolved_root) for path in files):
            return prepared_source
    return None if source_files_were_explicit else current_source


def _path_is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.relative_to(parent)
        return True
    except ValueError:
        return False


def _explicit_source_root(files: list[Path], raw_dir: Path) -> Path:
    resolved_raw_dir = raw_dir.resolve(strict=False)
    for path in files:
        resolved_path = path.resolve(strict=False)
        try:
            relative = resolved_path.relative_to(resolved_raw_dir)
        except ValueError:
            return path.parent
        if relative.parts:
            return raw_dir / relative.parts[0]
    return raw_dir


def _raw_dir_for_metadata(raw_dir: Path, metadata: dict | None, files: list[Path], source_files_were_explicit: bool) -> Path:
    if metadata and metadata.get("source_slug"):
        candidate = raw_dir / str(metadata["source_slug"])
        if candidate.exists():
            return candidate
    if source_files_were_explicit:
        return _explicit_source_root(files, raw_dir)
    return raw_dir


def clear_indexed_source_artifacts() -> None:
    if CHROMA_DIR.exists():
        CHROMA_DIR.mkdir(parents=True, exist_ok=True)
        try:
            client = chromadb.PersistentClient(path=str(CHROMA_DIR))
        except Exception:
            _reset_chroma_dir()
        else:
            try:
                client.delete_collection(CHROMA_COLLECTION_NAME)
            except Exception as exc:
                if _is_invalid_chroma_database_error(exc):
                    _reset_chroma_dir()
                elif not _is_missing_chroma_collection_error(exc):
                    raise
    if CURRENT_SOURCE_PATH.exists():
        CURRENT_SOURCE_PATH.unlink()


def _reset_chroma_dir() -> None:
    if CHROMA_DIR.exists():
        shutil.rmtree(CHROMA_DIR)
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)


def _is_missing_chroma_collection_error(exc: Exception) -> bool:
    not_found_error = getattr(getattr(chromadb, "errors", None), "NotFoundError", None)
    if isinstance(not_found_error, type) and isinstance(exc, not_found_error):
        return True
    return isinstance(exc, ValueError) and "does not exist" in str(exc).lower()


def _is_invalid_chroma_database_error(exc: Exception) -> bool:
    message = str(exc).lower()
    return (
        "no such table: tenants" in message
        or "attempt to write a readonly database" in message
        or "database disk image is malformed" in message
    )


def _documents_from_source_files(files: Iterable[Path]) -> list[Document]:
    documents: list[Document] = []
    for path in files:
        text = _safe_read_text(path).strip()
        if not re.search(r"\w", text):
            continue
        documents.append(
            Document(
                text=text,
                metadata={
                    "file_name": path.as_posix(),
                    "source": "repo_source",
                },
            )
        )
    return documents


def _split_documents(
    documents: Iterable[Document],
    chunk_size_tokens: int = CHUNK_SIZE_TOKENS,
    chunk_overlap_tokens: int = CHUNK_OVERLAP_TOKENS,
) -> list:
    splitter = SentenceSplitter(chunk_size=chunk_size_tokens, chunk_overlap=chunk_overlap_tokens)
    return splitter.get_nodes_from_documents(list(documents))


def run_chunking_ablation(
    source_files: Iterable[Path] | None = None,
    raw_dir: Path = RAW_DIR,
    chunk_sizes: Iterable[int] = CHUNK_ABLATION_SIZES,
    chunk_overlap_tokens: int = CHUNK_OVERLAP_TOKENS,
    output_path: Path | None = None,
) -> list[dict[str, float | int | str]]:
    """Compare chunk counts for multiple chunk sizes without rebuilding Chroma."""

    files = list(source_files) if source_files is not None else load_or_download_sources(raw_dir=raw_dir)
    documents = _documents_from_source_files(files)
    if not documents:
        raise RuntimeError("No readable text documents were found in the provided source files.")

    rows: list[dict[str, float | int | str]] = []
    for chunk_size in chunk_sizes:
        nodes = _split_documents(
            documents,
            chunk_size_tokens=int(chunk_size),
            chunk_overlap_tokens=min(int(chunk_overlap_tokens), max(int(chunk_size) - 1, 0)),
        )
        lengths = [len(node.get_content()) for node in nodes]
        rows.append(
            {
                "chunk_size": int(chunk_size),
                "chunk_overlap": int(min(chunk_overlap_tokens, max(int(chunk_size) - 1, 0))),
                "file_count": len(documents),
                "chunk_count": len(nodes),
                "avg_chunk_chars": round(sum(lengths) / len(lengths), 3) if lengths else 0.0,
                "max_chunk_chars": max(lengths) if lengths else 0,
            }
        )

    if output_path is not None:
        import pandas as pd

        output_path.parent.mkdir(parents=True, exist_ok=True)
        pd.DataFrame(rows).to_csv(output_path, index=False)
    return rows


def run_embedding_model_comparison(
    source_files: Iterable[Path] | None = None,
    raw_dir: Path = RAW_DIR,
    models: Iterable[str] = EMBEDDING_COMPARISON_MODELS,
    sample_size: int = 8,
    allow_model_downloads: bool | None = None,
    output_path: Path | None = None,
) -> list[dict[str, float | int | str]]:
    """Compare local embedding models on a small chunk sample without changing the active index."""

    files = list(source_files) if source_files is not None else load_or_download_sources(raw_dir=raw_dir)
    documents = _documents_from_source_files(files)
    if not documents:
        raise RuntimeError("No readable text documents were found in the provided source files.")
    nodes = _split_documents(documents)
    samples = [node.get_content() for node in nodes[: max(int(sample_size), 1)]]
    downloads_allowed = allow_model_downloads if allow_model_downloads is not None else os.getenv("ALLOW_MODEL_DOWNLOADS") == "1"

    rows: list[dict[str, float | int | str]] = []
    for model_name in models:
        started = time.perf_counter()
        if not downloads_allowed:
            rows.append(
                {
                    "model": str(model_name),
                    "status": "skipped_model_downloads_disabled",
                    "sample_count": len(samples),
                    "embedding_dim": 0,
                    "embedding_ms": 0.0,
                    "error": "",
                }
            )
            continue
        try:
            embed_model = HuggingFaceEmbedding(model_name=str(model_name))
            embeddings = embed_model.get_text_embedding_batch(samples)
            first_embedding = embeddings[0] if embeddings else []
            rows.append(
                {
                    "model": str(model_name),
                    "status": "ok",
                    "sample_count": len(samples),
                    "embedding_dim": len(first_embedding),
                    "embedding_ms": round((time.perf_counter() - started) * 1000, 3),
                    "error": "",
                }
            )
        except Exception as exc:
            rows.append(
                {
                    "model": str(model_name),
                    "status": "error",
                    "sample_count": len(samples),
                    "embedding_dim": 0,
                    "embedding_ms": round((time.perf_counter() - started) * 1000, 3),
                    "error": str(exc)[:200],
                }
            )

    if output_path is not None:
        import pandas as pd

        output_path.parent.mkdir(parents=True, exist_ok=True)
        pd.DataFrame(rows).to_csv(output_path, index=False)
    return rows


def build_index(
    source_files: Iterable[Path] | None = None,
    raw_dir: Path = RAW_DIR,
) -> tuple[VectorStoreIndex, list]:
    """Chunk source files, embed them locally, and persist vectors in ChromaDB."""

    import logging
    logger = logging.getLogger(__name__)

    source_files_were_explicit = source_files is not None
    current_source = load_current_source()
    current_raw_dir = _current_source_raw_dir(raw_dir, current_source)
    files = (
        list(source_files)
        if source_files_were_explicit
        else load_or_download_sources(raw_dir=current_raw_dir or raw_dir)
    )
    if not files:
        raise RuntimeError("No source documents found in data/raw and download produced no files.")
    logger.info("Building index from %d source files", len(files))

    documents = _documents_from_source_files(files)

    if not documents:
        raise RuntimeError("No readable text documents were found in the provided source files.")

    nodes = _split_documents(documents)

    try:
        for attempt in range(2):
            try:
                CHROMA_DIR.mkdir(parents=True, exist_ok=True)
                try:
                    chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
                except Exception:
                    _reset_chroma_dir()
                    chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
                try:
                    chroma_client.delete_collection(CHROMA_COLLECTION_NAME)
                except Exception as exc:
                    if _is_invalid_chroma_database_error(exc):
                        _reset_chroma_dir()
                        chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
                    elif not _is_missing_chroma_collection_error(exc):
                        raise
                try:
                    chroma_collection = chroma_client.create_collection(CHROMA_COLLECTION_NAME)
                except Exception as exc:
                    if not _is_invalid_chroma_database_error(exc):
                        raise
                    _reset_chroma_dir()
                    chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
                    chroma_collection = chroma_client.create_collection(CHROMA_COLLECTION_NAME)
                vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
                storage_context = StorageContext.from_defaults(vector_store=vector_store)
                embed_model = HuggingFaceEmbedding(model_name=EMBED_MODEL)

                start = time.perf_counter()
                index = VectorStoreIndex(nodes, storage_context=storage_context, embed_model=embed_model)
                embedding_time = time.perf_counter() - start
                break
            except Exception as exc:
                if attempt == 0 and _is_invalid_chroma_database_error(exc):
                    _reset_chroma_dir()
                    continue
                raise

        logger.info("Index built: %d documents, %d chunks, %.2fs embedding", len(documents), len(nodes), embedding_time)

        prepared_metadata = _metadata_for_files(files, current_source, source_files_were_explicit, raw_dir)
        active_raw_dir = _raw_dir_for_metadata(raw_dir, prepared_metadata, files, source_files_were_explicit)
        write_current_source(
            raw_dir=active_raw_dir,
            file_count=len(documents),
            chunk_count=len(nodes),
            source_input=str(prepared_metadata.get("source_input", "")) if prepared_metadata else "",
            source_type=str(prepared_metadata.get("source_type")) if prepared_metadata and prepared_metadata.get("source_type") else None,
            source_slug=str(prepared_metadata.get("source_slug")) if prepared_metadata and prepared_metadata.get("source_slug") else None,
        )
    except Exception:
        clear_indexed_source_artifacts()
        raise

    return index, nodes


def main() -> None:
    build_index()


if __name__ == "__main__":
    main()
